import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_color):
    """Sets background color for a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding in dxa (1 pt = 20 dxa)."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_code_block(doc, code_text):
    """Adds a styled code block with light background shading and monospace font."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, 'F1F5F9') # light slate gray background
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    
    # Add empty paragraph after table for spacing
    doc.add_paragraph()

def build_document():
    doc = Document()
    
    # Set Margins (1 inch all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x41, 0x55) # Slate 700
    
    # -------------------------------------------------------------
    # TITLE & HEADER
    # -------------------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(12)
    title_p.paragraph_format.space_after = Pt(4)
    title_run = title_p.add_run("Linear Algebra & Data Science Web Explorer")
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A) # Dark navy
    
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(24)
    sub_run = sub_p.add_run("Project Technical Implementation Report & Topic Explanations")
    sub_run.font.name = 'Calibri'
    sub_run.font.size = Pt(14)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    
    # Divider line table
    div_tbl = doc.add_table(rows=1, cols=1)
    div_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    div_cell = div_tbl.cell(0, 0)
    set_cell_background(div_cell, '2563EB') # Blue accent
    div_p = div_cell.paragraphs[0]
    div_p.paragraph_format.space_before = Pt(1)
    div_p.paragraph_format.space_after = Pt(1)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # -------------------------------------------------------------
    # EXECUTIVE SUMMARY
    # -------------------------------------------------------------
    h1 = doc.add_heading("1. Executive Summary & Overview", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(10)
    p.add_run(
        "This project is a full-stack educational web application designed to demonstrate, compute, and visualize "
        "six fundamental topics in Linear Algebra & Data Science. Developed using Python and Django, the system features a dedicated "
        "math engine that performs step-by-step symbolic matrix calculations, verifies field axioms, computes 3D vector geometry, "
        "and presents full LaTeX math rendering directly on the web interface."
    )

    # Key Features Table
    table = doc.add_table(rows=7, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    headers = ["Unit / Topic", "Mathematical Concept", "Primary Solver Function"]
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], '1E293B')
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=120, right=120)
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    data = [
        ("Unit 1 • Topic 1", "Systems of Linear Equations & Gaussian Elimination", "solve_gaussian_elimination()"),
        ("Unit 1 • Topic 2", "Field Axioms via GF(2) (Galois Field of 2)", "analyze_gf2_field()"),
        ("Unit 1 • Topic 3", "3D Vector Products, Angles & Projections", "compute_vector_operations()"),
        ("Unit 2 • Topic 1", "Gram-Schmidt Orthogonalization Process", "solve_gram_schmidt()"),
        ("Unit 2 • Topic 2", "Cofactor Expansion for Determinants", "solve_cofactor_expansion()"),
        ("Unit 2 • Topic 3", "Eigenvalues, Eigenvectors & Matrix Diagonalization", "solve_diagonalization()"),
    ]
    
    for row_idx, row_data in enumerate(data, start=1):
        row_cells = table.rows[row_idx].cells
        bg_color = 'F8FAFC' if row_idx % 2 == 0 else 'FFFFFF'
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = text
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=120, right=120)
            if col_idx == 2:
                row_cells[col_idx].paragraphs[0].runs[0].font.name = 'Consolas'
                row_cells[col_idx].paragraphs[0].runs[0].font.size = Pt(9.5)
                row_cells[col_idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # -------------------------------------------------------------
    # SOFTWARE ARCHITECTURE
    # -------------------------------------------------------------
    h1 = doc.add_heading("2. High-Level Software Architecture", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.add_run(
        "The project strictly follows software engineering separation of concerns by isolating pure mathematical logic "
        "from HTTP request handling and template rendering:"
    )
    
    bullets = [
        ("Math Engine (math_engine.py): ", "Contains pure Python functions that implement algorithms from first principles. It uses SymPy for exact symbolic math and LaTeX generation, and NumPy for high-performance vectorized 3D geometry."),
        ("Controller Layer (views.py): ", "Handles HTTP GET and POST requests, parses matrix strings from textareas into float/symbol lists, passes data to the math engine, and injects results into template context."),
        ("Form Validation (forms.py): ", "Ensures user input is valid before attempting calculations, setting default fallback matrices when pages load for the first time."),
        ("Frontend Templates (HTML + MathJax): ", "Displays formatted mathematical solutions with step-by-step collapsible cards, formatted matrix boxes, and live MathJax typesetting."),
        ("Live Code Introspection: ", "Uses Python's inspect.getsource() module in views.py to extract and display the backend algorithm code directly on the web page for educational learning.")
    ]
    
    for b_title, b_desc in bullets:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(4)
        run_t = bp.add_run(b_title)
        run_t.bold = True
        run_t.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        bp.add_run(b_desc)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # -------------------------------------------------------------
    # TOPIC BY TOPIC DETAILED BREAKDOWN
    # -------------------------------------------------------------
    h1 = doc.add_heading("3. Detailed Implementation of the 6 Topics", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    # --- TOPIC 1 ---
    h2 = doc.add_heading("Topic 1: Systems of Linear Equations & Gaussian Elimination", level=2)
    h2.runs[0].font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.add_run("• Mathematical Principle: ").bold = True
    p.add_run("Solves linear systems [A | b] by executing Elementary Row Operations (EROs) to transform the augmented matrix into Reduced Row Echelon Form (RREF). System consistency is verified using the Rouché–Capelli Theorem:\n")
    p.add_run("  - Inconsistent (No Solution): rank(A) < rank([A|b])\n")
    p.add_run("  - Unique Solution: rank(A) = rank([A|b]) = n (number of variables)\n")
    p.add_run("  - Infinitely Many Solutions: rank(A) = rank([A|b]) < n (n - rank(A) free variables)\n")

    p = doc.add_paragraph()
    p.add_run("• Python Implementation Logic:").bold = True
    add_code_block(doc, 
"""def solve_gaussian_elimination(matrix_data):
    mat = sp.Matrix(matrix_data)
    rows, cols = mat.shape
    num_vars = cols - 1
    steps = []

    # 1. Iterative Row Reduction & Pivoting
    for c in range(num_vars):
        # Find pivot element & Swap rows if needed
        # Scale pivot row to 1: R_i <-- (1 / pivot) * R_i
        # Eliminate entries above and below pivot: R_k <-- R_k - (factor) * R_pivot
        # Save LaTeX snapshot of matrix after each ERO step

    # 2. Solution Classification via Rouché-Capelli Theorem
    rank_A = A_part.rank()
    rank_aug = mat.rank()
    if rank_A < rank_aug:
        solution_type = "Inconsistent System (No Solution)"
    elif rank_A == rank_aug == num_vars:
        solution_type = "Unique Solution"
    else:
        solution_type = "Infinitely Many Solutions (Parametric)"
""")

    # --- TOPIC 2 ---
    h2 = doc.add_heading("Topic 2: Field Axioms via GF(2)", level=2)
    h2.runs[0].font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.add_run("• Mathematical Principle: ").bold = True
    p.add_run("Evaluates the Galois Field F_2 = {0, 1} under modulo-2 addition (XOR) and modulo-2 multiplication (AND). Programmatically tests and proves all 11 algebraic field axioms (Closure, Associativity, Commutativity, Identities, Inverses, and Distributivity).\n")

    p = doc.add_paragraph()
    p.add_run("• Python Implementation Logic:").bold = True
    add_code_block(doc, 
"""def analyze_gf2_field():
    elements = [0, 1]
    add_table = [[(a + b) % 2 for b in elements] for a in elements] # Addition (XOR)
    mul_table = [[(a * b) % 2 for b in elements] for a in elements] # Multiplication (AND)

    # Explicitly check all 11 axioms across elements
    add_closed = all((a + b) % 2 in elements for a in elements for b in elements)
    add_assoc  = all(((a + b) + c) % 2 == (a + (b + c)) % 2 for a in elements for b in elements for c in elements)
    distrib    = all((a * ((b + c) % 2)) % 2 == ((a * b) + (a * c)) % 2 for a in elements for b in elements for c in elements)
    
    return {'add_table': add_table, 'mul_table': mul_table, 'axioms': axioms}
""")

    # --- TOPIC 3 ---
    h2 = doc.add_heading("Topic 3: 3D Vector Dot Product, Cross Product & Projections", level=2)
    h2.runs[0].font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.add_run("• Mathematical Principle: ").bold = True
    p.add_run("Computes vector operations for v1, v2 in R^3. Calculates norms ||v||, dot product v1 . v2, cross product v1 x v2, angle theta = arccos((v1 . v2)/(||v1|| ||v2||)), vector projection proj_v2(v1) = ((v1 . v2)/||v2||^2) * v2, and parallelogram area ||v1 x v2||.\n")

    p = doc.add_paragraph()
    p.add_run("• Python Implementation Logic:").bold = True
    add_code_block(doc, 
"""def compute_vector_operations(v1_list, v2_list):
    v1, v2 = np.array(v1_list), np.array(v2_list)
    mag1, mag2 = np.linalg.norm(v1), np.linalg.norm(v2)
    
    dot_prod = np.dot(v1, v2)
    cross_prod = np.cross(v1, v2)
    cross_mag = np.linalg.norm(cross_prod)
    
    # Angle calculation
    cos_theta = np.clip(dot_prod / (mag1 * mag2), -1.0, 1.0)
    angle_rad = np.arccos(cos_theta)
    
    # Vector Projection proj_v2(v1)
    proj_vector = (dot_prod / (mag2 ** 2)) * v2
""")

    # --- TOPIC 4 ---
    h2 = doc.add_heading("Topic 4: Gram-Schmidt Orthogonalization Process", level=2)
    h2.runs[0].font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.add_run("• Mathematical Principle: ").bold = True
    p.add_run("Transforms a set of linearly independent vectors {v_1, ..., v_k} into an orthogonal basis {u_1, ..., u_k} and an orthonormal basis {e_1, ..., e_k} using the Gram-Schmidt formula:\n")
    p.add_run("  u_i = v_i - sum_{j=1}^{i-1} proj_{u_j}(v_i),    e_i = u_i / ||u_i||\n")

    p = doc.add_paragraph()
    p.add_run("• Python Implementation Logic:").bold = True
    add_code_block(doc, 
"""def solve_gram_schmidt(vectors_data):
    orig_vectors = [sp.Matrix(v) for v in vectors_data]
    u_vectors, e_vectors = [], []

    for i in range(len(orig_vectors)):
        v_i = orig_vectors[i]
        u_i = v_i.copy()
        for j in range(i):
            u_j = u_vectors[j]
            proj_vec = (v_i.dot(u_j) / u_j.dot(u_j)) * u_j
            u_i = u_i - proj_vec
            
        u_vectors.append(u_i)
        u_norm = sp.sqrt(u_i.dot(u_i))
        e_vectors.append(u_i / u_norm if u_norm != 0 else u_i)
""")

    # --- TOPIC 5 ---
    h2 = doc.add_heading("Topic 5: Cofactor Expansion for Determinants", level=2)
    h2.runs[0].font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.add_run("• Mathematical Principle: ").bold = True
    p.add_run("Calculates det(A) for a square matrix A by performing Cofactor Expansion along any chosen row or column. Uses minor submatrices M_{r,c} and sign coefficients (-1)^{r+c} where Cofactor C_{r,c} = (-1)^{r+c} det(M_{r,c}).\n")

    p = doc.add_paragraph()
    p.add_run("• Python Implementation Logic:").bold = True
    add_code_block(doc, 
"""def solve_cofactor_expansion(matrix_data, expand_by='row', idx=0):
    mat = sp.Matrix(matrix_data)
    terms = []
    total_det = 0
    
    for i in range(n):
        val = mat[idx, i] if expand_by == 'row' else mat[i, idx]
        sub_mat = mat.minor_submatrix(r, c)
        sign = (-1) ** (r + c)
        minor_det = sub_mat.det()
        cofactor = sign * minor_det
        total_det += val * cofactor
        
        terms.append({'entry': val, 'submatrix': sub_mat, 'cofactor': cofactor})
""")

    # --- TOPIC 6 ---
    h2 = doc.add_heading("Topic 6: Eigenvalues, Eigenvectors & Diagonalization", level=2)
    h2.runs[0].font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.add_run("• Mathematical Principle: ").bold = True
    p.add_run("Computes the characteristic polynomial det(A - lambda I) = 0, finds eigenvalues lambda and eigenvectors v, evaluates algebraic vs geometric multiplicities, and constructs matrix factorization A = P D P^-1 if diagonalizability criteria are satisfied.\n")

    p = doc.add_paragraph()
    p.add_run("• Python Implementation Logic:").bold = True
    add_code_block(doc, 
"""def solve_diagonalization(matrix_data):
    mat = sp.Matrix(matrix_data)
    lam = sp.Symbol('\\lambda')
    char_poly = (mat - lam * sp.eye(n)).det()
    
    eigen_info = mat.eigenvects() # Returns (val, alg_mult, [vects])
    # Verify if sum(geometric_multiplicities) == n
    if total_geometric_mult == n:
        P_mat = sp.Matrix.hstack(*P_cols) # Matrix of eigenvectors
        D_mat = sp.diag(*D_diag)          # Diagonal matrix of eigenvalues
        verification = P_mat * D_mat * P_mat.inv() # Proves A = P D P^-1
""")

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # -------------------------------------------------------------
    # CONCLUSION & SKILLS LEARNED
    # -------------------------------------------------------------
    h1 = doc.add_heading("4. Summary of Developer Skills & Learnings", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.add_run("Through developing this project, key software engineering and mathematical computer science skills were mastered:\n")
    
    skills = [
        ("Exact Symbolic Computation: ", "Leveraged SymPy to eliminate round-off floating-point errors, ensuring exact fractional and algebraic math outputs."),
        ("Full-Stack Django Integration: ", "Designed clean MVT request cycles, input sanitization pipelines, and template context handling."),
        ("Frontend Mathematical Typesetting: ", "Mastered dynamic MathJax integration to render textbook LaTeX notation seamlessly in web browsers."),
        ("Modular Architecture: ", "Decoupled pure algorithms (math_engine.py) from web framework views, facilitating unit testing and code maintainability.")
    ]
    
    for s_title, s_desc in skills:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(4)
        run_t = bp.add_run(s_title)
        run_t.bold = True
        run_t.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        bp.add_run(s_desc)
        
    doc.save("d:/Linear Algebra Website/Linear_Algebra_Project_Explanation.docx")
    print("Document successfully created at d:/Linear Algebra Website/Linear_Algebra_Project_Explanation.docx")

if __name__ == '__main__':
    build_document()
